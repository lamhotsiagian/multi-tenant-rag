"""
Document processing service.

Responsibilities:
  - File upload validation and storage.
  - Text extraction (PDF / DOCX / TXT).
  - Chunking, embedding, and vector indexing.
  - Document lifecycle management (list, get, delete).

Design principles:
  - This service raises *domain exceptions* (``DocumentNotFoundError``,
    ``DocumentProcessingError``, ``DocumentValidationError``).
    HTTP mapping is the router's responsibility.
  - ``EmbeddingService`` and ``QdrantVectorService`` are injected via the
    constructor — the service never instantiates them itself.
  - File I/O uses ``aiofiles`` for non-blocking reads/writes.
  - ``chunk_record.vector_id`` is assigned *before* the ORM object is added
    to the session to ensure the UUID is deterministic.
"""
import os
import uuid
import aiofiles
import structlog
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import UploadFile
from sqlalchemy import and_
from sqlalchemy.orm import Session

import PyPDF2
from docx import Document as DocxDocument

from app.models.document import Document, DocumentChunk
from app.models.tenant import Tenant
from app.services.embedding_service import EmbeddingService
from app.services.vector_service import QdrantVectorService
from app.config import settings

logger = structlog.get_logger(__name__)


# ── Domain exceptions ─────────────────────────────────────────────────────────

class DocumentNotFoundError(Exception):
    """Raised when a document cannot be located for the given tenant."""


class DocumentValidationError(Exception):
    """Raised when an uploaded file fails validation."""


class DocumentProcessingError(Exception):
    """Raised when text extraction or embedding fails."""


# ── Service ───────────────────────────────────────────────────────────────────

class DocumentService:
    """
    Manages the full document lifecycle within a tenant.

    Args:
        embedding_service: Injected embedding singleton.
        vector_service: Injected Qdrant singleton.
    """

    def __init__(
        self,
        embedding_service: EmbeddingService,
        vector_service: QdrantVectorService,
    ) -> None:
        self.embedding_service = embedding_service
        self.vector_service = vector_service
        self.upload_dir = Path(settings.upload_dir)
        self.max_file_size = settings.max_file_size_mb * 1024 * 1024
        self.allowed_types: List[str] = list(settings.allowed_file_types)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    # ── Validation ────────────────────────────────────────────────────────────

    def _validate_file(self, file: UploadFile) -> None:
        """
        Validate file size and extension.

        Raises:
            DocumentValidationError: for oversized or disallowed file types.
        """
        if hasattr(file, "size") and file.size and file.size > self.max_file_size:
            raise DocumentValidationError(
                f"File size exceeds the {settings.max_file_size_mb} MB limit"
            )

        if file.filename:
            ext = file.filename.rsplit(".", 1)[-1].lower()
            if ext not in self.allowed_types:
                raise DocumentValidationError(
                    f"File type '.{ext}' is not allowed. "
                    f"Accepted types: {', '.join(self.allowed_types)}"
                )

    # ── Upload ────────────────────────────────────────────────────────────────

    async def upload_document(
        self,
        db: Session,
        tenant_id: str,
        file: UploadFile,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Document:
        """
        Persist an uploaded file to disk and create its database record.

        Args:
            db: SQLAlchemy session.
            tenant_id: Owning tenant UUID string.
            file: FastAPI ``UploadFile`` object.
            metadata: Optional dict of extra document metadata.

        Returns:
            Newly created ``Document`` ORM instance (status=``"uploaded"``).

        Raises:
            DocumentValidationError: if the file fails validation.
            DocumentProcessingError: if storage or database write fails.
        """
        self._validate_file(file)

        file_id = str(uuid.uuid4())
        ext = file.filename.rsplit(".", 1)[-1].lower() if file.filename else "txt"
        stored_filename = f"{tenant_id}_{file_id}.{ext}"
        file_path = self.upload_dir / stored_filename

        try:
            content = await file.read()
            async with aiofiles.open(file_path, "wb") as fh:
                await fh.write(content)

            document = Document(
                tenant_id=tenant_id,
                filename=stored_filename,
                original_filename=file.filename or "unknown",
                content_type=file.content_type or "application/octet-stream",
                file_size=len(content),
                file_path=str(file_path),
                status="uploaded",
                doc_metadata=metadata or {},
            )
            db.add(document)
            db.commit()
            db.refresh(document)

            logger.info("Document uploaded", document_id=str(document.id), tenant_id=tenant_id)
            return document

        except DocumentValidationError:
            raise
        except Exception as exc:
            if file_path.exists():
                file_path.unlink(missing_ok=True)
            logger.error("Document upload failed", error=str(exc))
            raise DocumentProcessingError(f"Failed to store document: {exc}") from exc

    # ── Text extraction ───────────────────────────────────────────────────────

    def extract_text_from_file(self, file_path: str, content_type: str) -> str:
        """
        Dispatch text extraction to the appropriate handler based on MIME type
        or file extension.

        Returns an empty string on extraction failure rather than propagating
        — callers should treat an empty result as a failed document.
        """
        try:
            if content_type == "application/pdf" or file_path.endswith(".pdf"):
                return self._extract_from_pdf(file_path)
            elif (
                content_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                or file_path.endswith(".docx")
            ):
                return self._extract_from_docx(file_path)
            else:
                # Treat everything else as plain text
                return self._extract_from_txt(file_path)
        except Exception as exc:
            logger.error("Text extraction failed", file=file_path, error=str(exc))
            return ""

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text page-by-page from a PDF, prefixing each page number."""
        pages: List[str] = []
        with open(file_path, "rb") as fh:
            reader = PyPDF2.PdfReader(fh)
            for page_num, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages.append(f"[Page {page_num}]\n{page_text}")
        return "\n\n".join(pages)

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract paragraph text from a DOCX document."""
        doc = DocxDocument(file_path)
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def _extract_from_txt(self, file_path: str) -> str:
        """Read a plain-text file, ignoring un-decodable bytes."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as fh:
            return fh.read()

    # ── Processing pipeline ───────────────────────────────────────────────────

    async def process_document(
        self,
        db: Session,
        document_id: str,
        tenant_id: str,
    ) -> bool:
        """
        Full processing pipeline: extract → chunk → embed → index.

        Status transitions: ``uploaded`` → ``processing`` → ``processed | failed``.

        Returns:
            ``True`` on success, ``False`` on failure (status already set to
            ``"failed"`` in the database before returning).
        """
        document = self.get_document(db, document_id, tenant_id)
        if not document:
            logger.error("Document not found", document_id=document_id)
            return False

        try:
            document.status = "processing"
            db.commit()

            # 1. Extract raw text
            text_content = self.extract_text_from_file(
                document.file_path, document.content_type
            )
            if not text_content.strip():
                logger.error("No text extracted", document_id=document_id)
                document.status = "failed"
                db.commit()
                return False

            document.word_count = len(text_content.split())

            # 2. Chunk
            chunks = self.embedding_service.chunk_text_for_embedding(
                text_content, max_chunk_size=512, overlap_size=50
            )

            # 3. Embed
            embedded_chunks = await self.embedding_service.embed_document_chunks(chunks)

            # 4. Persist chunks to DB and vector store
            chunk_records: List[DocumentChunk] = []
            vector_documents: List[Dict[str, Any]] = []

            for chunk_data in embedded_chunks:
                # Assign vector_id BEFORE creating the ORM object so it is
                # deterministic and can be referenced without a DB flush.
                vector_id = str(uuid.uuid4())

                chunk_record = DocumentChunk(
                    document_id=document.id,
                    tenant_id=tenant_id,
                    chunk_index=chunk_data["chunk_index"],
                    text_content=chunk_data["text"],
                    chunk_size=chunk_data["chunk_size"],
                    start_char=chunk_data.get("start_char"),
                    end_char=chunk_data.get("end_char"),
                    vector_id=vector_id,
                    embedding_model=chunk_data["embedding_model"],
                    embedding_dimension=chunk_data["embedding_dimension"],
                )
                chunk_records.append(chunk_record)

                vector_documents.append(
                    {
                        "document_id": str(document.id),
                        "chunk_id": vector_id,
                        "text": chunk_data["text"],
                        "embedding": chunk_data["embedding"],
                        "source": document.original_filename,
                        "page_number": chunk_data.get("page_number"),
                        "metadata": {
                            "filename": document.original_filename,
                            "content_type": document.content_type,
                            "chunk_index": chunk_data["chunk_index"],
                            "start_char": chunk_data.get("start_char"),
                            "end_char": chunk_data.get("end_char"),
                        },
                    }
                )

            db.add_all(chunk_records)

            success = await self.vector_service.add_documents(
                tenant_id=tenant_id,
                documents=vector_documents,
            )

            if success:
                document.status = "processed"
                document.total_chunks = len(chunks)
                document.processed_chunks = len(chunks)
                document.collection_name = self.vector_service.default_collection
                if embedded_chunks:
                    document.embedding_model = embedded_chunks[0]["embedding_model"]
                db.commit()
                logger.info(
                    "Document processed",
                    document_id=document_id,
                    chunks=len(chunks),
                )
                return True
            else:
                document.status = "failed"
                db.commit()
                return False

        except Exception as exc:
            logger.error("Document processing error", document_id=document_id, error=str(exc))
            document.status = "failed"
            db.commit()
            return False

    # ── CRUD helpers ──────────────────────────────────────────────────────────

    async def delete_document(
        self,
        db: Session,
        document_id: str,
        tenant_id: str,
    ) -> bool:
        """
        Delete a document's vectors, file, and database record.

        Returns:
            ``True`` on success, ``False`` when the document was not found.
        """
        document = self.get_document(db, document_id, tenant_id)
        if not document:
            return False

        try:
            await self.vector_service.delete_document(tenant_id, document_id)

            if os.path.exists(document.file_path):
                os.remove(document.file_path)

            db.delete(document)
            db.commit()
            logger.info("Document deleted", document_id=document_id)
            return True
        except Exception as exc:
            logger.error("Document deletion failed", document_id=document_id, error=str(exc))
            return False

    def get_document(
        self,
        db: Session,
        document_id: str,
        tenant_id: str,
    ) -> Optional[Document]:
        """Return a single document scoped to the given tenant, or ``None``."""
        return (
            db.query(Document)
            .filter(
                and_(
                    Document.id == document_id,
                    Document.tenant_id == tenant_id,
                )
            )
            .first()
        )

    def list_documents(
        self,
        db: Session,
        tenant_id: str,
        skip: int = 0,
        limit: int = 100,
        status_filter: Optional[str] = None,
    ) -> List[Document]:
        """
        Return a paginated list of documents for a tenant.

        Args:
            db: SQLAlchemy session.
            tenant_id: UUID string of the owning tenant.
            skip: Number of records to skip (for pagination).
            limit: Maximum number of records to return.
            status_filter: Optional status string to filter by.
        """
        query = db.query(Document).filter(Document.tenant_id == tenant_id)
        if status_filter:
            query = query.filter(Document.status == status_filter)
        return query.order_by(Document.created_at.desc()).offset(skip).limit(limit).all()